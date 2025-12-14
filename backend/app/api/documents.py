# app/api/documents.py
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
import os
from pypdf import PdfReader
from sqlalchemy.orm import Session
from typing import List
import pandas as pd
from docx import Document as DocxDocument
import openpyxl

from app.core.db import get_db
from app.models.db_models import Document, Chunk
from app.models.user import User
from app.core.rag import add_chunks
from app.core.llm_client import LLMClient
from app.core.security import get_current_user

router = APIRouter(tags=["documents"])

UPLOAD_DIR = "./uploaded_docs"
os.makedirs(UPLOAD_DIR, exist_ok=True)


# ==================== FILE TYPE CHECKERS ====================

def is_pdf_file(filename: str, content_type: str) -> bool:
    """Check if file is PDF"""
    return filename.lower().endswith(".pdf") or content_type == "application/pdf"


def is_image_file(filename: str, content_type: str) -> bool:
    """Check if file is image"""
    image_extensions = {".png", ".jpg", ".jpeg"}
    image_content_types = {"image/png", "image/jpeg"}
    ext = os.path.splitext(filename.lower())[1]
    return ext in image_extensions or content_type in image_content_types


def is_excel_file(filename: str, content_type: str) -> bool:
    """Check if file is Excel"""
    excel_extensions = {".xlsx", ".xls"}
    excel_content_types = {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel"
    }
    ext = os.path.splitext(filename.lower())[1]
    return ext in excel_extensions or content_type in excel_content_types


def is_word_file(filename: str, content_type: str) -> bool:
    """Check if file is Word"""
    word_extensions = {".docx", ".doc"}
    word_content_types = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword"
    }
    ext = os.path.splitext(filename.lower())[1]
    return ext in word_extensions or content_type in word_content_types


def is_text_file(filename: str, content_type: str) -> bool:
    """Check if file is text or CSV"""
    text_extensions = {".txt", ".csv"}
    text_content_types = {"text/plain", "text/csv"}
    ext = os.path.splitext(filename.lower())[1]
    return ext in text_extensions or content_type in text_content_types


# ==================== EXTRACTORS ====================

def extract_pdf_text(path: str) -> List[str]:
    """Extract text from PDF, page by page"""
    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return pages


def extract_excel_text(path: str) -> List[str]:
    """Extract text from Excel, sheet by sheet"""
    try:
        # Read all sheets
        xl_file = pd.ExcelFile(path)
        chunks = []
        
        for sheet_name in xl_file.sheet_names:
            df = pd.read_excel(path, sheet_name=sheet_name)
            
            # Convert dataframe to readable text
            text_lines = [f"Sheet: {sheet_name}\n"]
            text_lines.append(df.to_string(index=False))
            
            chunk_text = "\n".join(text_lines)
            if chunk_text.strip():
                chunks.append(chunk_text)
        
        return chunks
    except Exception as e:
        print(f"Error extracting Excel: {e}")
        return [f"Error reading Excel file: {str(e)}"]


def extract_word_text(path: str) -> List[str]:
    """Extract text from Word document, paragraph by paragraph"""
    try:
        doc = DocxDocument(path)
        
        # Combine all paragraphs into one chunk
        # (You can split by sections if needed)
        text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        if text.strip():
            return [text]
        return ["Empty document"]
    except Exception as e:
        print(f"Error extracting Word: {e}")
        return [f"Error reading Word file: {str(e)}"]


def extract_text_file(path: str) -> List[str]:
    """Extract text from TXT or CSV"""
    try:
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        if content.strip():
            return [content]
        return ["Empty file"]
    except Exception as e:
        print(f"Error reading text file: {e}")
        return [f"Error reading file: {str(e)}"]


# ==================== UPLOAD ENDPOINT ====================

@router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # Determine file type
    is_pdf = is_pdf_file(file.filename, file.content_type or "")
    is_image = is_image_file(file.filename, file.content_type or "")
    is_excel = is_excel_file(file.filename, file.content_type or "")
    is_word = is_word_file(file.filename, file.content_type or "")
    is_text = is_text_file(file.filename, file.content_type or "")
    
    # Check if file type is supported
    if not any([is_pdf, is_image, is_excel, is_word, is_text]):
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Supported: PDF, DOCX, XLSX, XLS, CSV, TXT, PNG, JPG, JPEG"
        )
    
    # Save file locally
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())

    # 🔒 Insert into documents table with user_id
    doc = Document(
        name=file.filename,
        path=file_path,
        user_id=current_user.id
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    chunk_texts = []
    metadatas = []

    # Extract based on file type
    if is_pdf:
        pages = extract_pdf_text(file_path)
        chunk_texts = [page for page in pages if page.strip()]
        metadatas = [
            {
                "document_id": doc.id,
                "page": idx,
                "source": "pdf",
                "user_id": current_user.id
            }
            for idx in range(len(chunk_texts))
        ]
    
    elif is_excel:
        sheets = extract_excel_text(file_path)
        chunk_texts = [sheet for sheet in sheets if sheet.strip()]
        metadatas = [
            {
                "document_id": doc.id,
                "page": idx,
                "source": "excel",
                "filename": file.filename,
                "user_id": current_user.id
            }
            for idx in range(len(chunk_texts))
        ]
    
    elif is_word:
        paragraphs = extract_word_text(file_path)
        chunk_texts = [para for para in paragraphs if para.strip()]
        metadatas = [
            {
                "document_id": doc.id,
                "page": 0,
                "source": "word",
                "filename": file.filename,
                "user_id": current_user.id
            }
        ]
    
    elif is_text:
        content = extract_text_file(file_path)
        chunk_texts = [c for c in content if c.strip()]
        metadatas = [
            {
                "document_id": doc.id,
                "page": 0,
                "source": "text",
                "filename": file.filename,
                "user_id": current_user.id
            }
        ]
    
    elif is_image:
        # Handle image files with Gemini Vision
        llm_client = LLMClient()
        text = llm_client.extract_image_text(file_path)
        
        if text.strip():
            chunk_texts = [text]
            metadatas = [
                {
                    "document_id": doc.id,
                    "page": 0,
                    "source": "image",
                    "filename": file.filename,
                    "user_id": current_user.id
                }
            ]

    # Insert chunks into Chroma + SQLite
    if chunk_texts:
        add_chunks(chunk_texts, metadatas)

        # Store chunk rows in SQLite
        for idx, text in enumerate(chunk_texts):
            meta_info = {"page": metadatas[idx].get("page", idx)}
            meta_info["source"] = metadatas[idx].get("source", "unknown")
            if "filename" in metadatas[idx]:
                meta_info["filename"] = metadatas[idx]["filename"]
            
            chunk = Chunk(
                document_id=doc.id,
                content=text,
                meta_json=str(meta_info),
                user_id=current_user.id
            )
            db.add(chunk)

    db.commit()
    
    return {
        "message": "uploaded",
        "document_id": doc.id,
        "chunks": len(chunk_texts),
        "file_type": metadatas[0].get("source", "unknown") if metadatas else "unknown"
    }